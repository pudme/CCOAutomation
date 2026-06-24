from __future__ import annotations

import asyncio
import csv
import difflib
import hashlib
import io
import json
import re
import uuid
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Awaitable, Callable

import chromadb
from loguru import logger
from minio import Minio
from sqlalchemy import func, or_, select
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import selectinload

from ai.tools import ingest_text
from config import get_settings
from database import AsyncSessionLocal
from models.compliance import (
    AppSetting,
    Control,
    ControlStatus,
    DataImport,
    EvidenceItem,
    EvidenceStatus,
    EvidenceType,
    ImportStatus,
    PersonnelRecord,
)
from ai.gateway import MODEL_HAIKU, call_claude, is_daily_limit_exception
from services.change_log import log_change

settings = get_settings()

TYPE_CHOICES = {
    "active_employee_list",
    "terminated_employee_list",
    "mfa_enrollment",
    "training_completion",
    "entra_id_export",
    "intune_compliance",
    "crowdstrike_inventory",
    "access_rights_report",
    "evidence_document",
    "policy_document",
    "audit_log",
    "auditor_checklist",
    "unknown",
}
PROCESS_TIMEOUT_SECONDS = 300
MAX_IMPORT_RETRIES = 2
_RAW_MEETING_NOTE_TYPES = {
    "meeting_note",
    "meeting_notes",
    "meeting_minutes",
    "meeting_minute",
    "raw_meeting_notes",
    "raw_notes",
}
_CONTROL_ID_PATTERNS = [
    r"\b(Cl\.\d+\.\d+)\b",
    r"\b(A\.\d+\.\d+)\b",
    r"\b([A-Z]{2,4}\.L[12]-[\d\.]+)\b",
    r"\b([A-Z]{2,4}\.\d+\.\d{3})\b",
]
_DPA_REQUEST_ID_PATTERN = re.compile(r"\bDPA[-_\s]*0*([1-9]|1[0-9]|2[0-3])\b", re.IGNORECASE)
_REANALYZE_CHECKPOINT_KEY = "reanalyze_checkpoint"
_BLOCKED_IMPORT_FILENAMES = {"readme.md"}
_DUPLICATE_TOKEN_STOPWORDS = {
    "the",
    "and",
    "for",
    "with",
    "from",
    "this",
    "that",
    "document",
    "evidence",
    "report",
    "file",
    "export",
    "data",
}
_IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp", ".tif", ".tiff"}
_SPREADSHEET_EXTENSIONS = {".csv", ".xlsx", ".xls", ".xlsm", ".ods"}
_DOCUMENT_EXTENSIONS = {".pdf", ".doc", ".docx", ".ppt", ".pptx", ".txt", ".md", ".rtf"}
_STRUCTURED_DATA_EXTENSIONS = {".json", ".yaml", ".yml", ".xml"}


def _is_blocked_import_filename(filename: str) -> bool:
    normalized = Path(filename or "").name.strip().lower()
    if not normalized:
        return False
    return normalized in _BLOCKED_IMPORT_FILENAMES


async def _get_reanalyze_checkpoint(session: AsyncSession) -> int:
    row = (
        await session.execute(
            select(AppSetting).where(AppSetting.key == _REANALYZE_CHECKPOINT_KEY)
        )
    ).scalars().first()
    if row is None:
        return 0
    try:
        return int(str(row.value or "0").strip())
    except ValueError:
        return 0


async def _set_reanalyze_checkpoint(session: AsyncSession, import_id: int) -> None:
    row = (
        await session.execute(
            select(AppSetting).where(AppSetting.key == _REANALYZE_CHECKPOINT_KEY)
        )
    ).scalars().first()
    now = datetime.now(timezone.utc).isoformat()
    if row is None:
        row = AppSetting(
            key=_REANALYZE_CHECKPOINT_KEY,
            value=str(import_id),
            updated_at=now,
        )
        session.add(row)
    else:
        row.value = str(import_id)
        row.updated_at = now
    await session.commit()


async def _clear_reanalyze_checkpoint(session: AsyncSession) -> None:
    row = (
        await session.execute(
            select(AppSetting).where(AppSetting.key == _REANALYZE_CHECKPOINT_KEY)
        )
    ).scalars().first()
    if row is None:
        return
    await session.delete(row)
    await session.commit()


def get_minio_client() -> Minio:
    return Minio(
        endpoint=settings.minio_endpoint,
        access_key=settings.minio_access_key,
        secret_key=settings.minio_secret_key,
        secure=False,
    )


def build_import_object_name(filename: str) -> str:
    safe_name = Path(filename).name
    return f"imports/{uuid.uuid4()}_{safe_name}"


def _embed_chunks_to_collection(
    *,
    client: chromadb.HttpClient,
    collection_name: str,
    import_id: int,
    chunks: list[str],
    safe_metadata: dict[str, str],
) -> None:
    collection = client.get_or_create_collection(collection_name)
    try:
        collection.delete(where={"import_id": str(import_id)})
    except Exception:
        pass
    chroma_batch_size = 5000
    for start in range(0, len(chunks), chroma_batch_size):
        end = min(start + chroma_batch_size, len(chunks))
        batch_docs = chunks[start:end]
        ids = [f"import-{import_id}-{index}" for index in range(start, end)]
        metadatas = []
        for index in range(start, end):
            chunk_meta = dict(safe_metadata)
            chunk_meta["chunk_index"] = str(index)
            metadatas.append(chunk_meta)
        try:
            collection.upsert(ids=ids, documents=batch_docs, metadatas=metadatas)
        except Exception:
            collection.add(ids=ids, documents=batch_docs, metadatas=metadatas)


def _embedding_collections_for_type(detected_type: str | None) -> tuple[str, ...]:
    normalized = str(detected_type or "").strip().lower()
    if normalized in _RAW_MEETING_NOTE_TYPES:
        return ("meeting_notes",)
    return ("compliance_docs",)


def embed_import_text(
    import_id: int,
    text: str,
    metadata: dict[str, str],
    *,
    detected_type: str | None = None,
    force_collections: tuple[str, ...] | None = None,
) -> None:
    if not text.strip():
        return
    client = chromadb.HttpClient(host=settings.chroma_host, port=settings.chroma_port)
    safe_metadata = {str(key): str(value) for key, value in metadata.items()}
    safe_metadata["import_id"] = str(import_id)
    if detected_type:
        safe_metadata["detected_type"] = str(detected_type)
    chunk_size = 4000
    chunks = [text[idx : idx + chunk_size] for idx in range(0, len(text), chunk_size)] or [text]
    target_collections = force_collections or _embedding_collections_for_type(detected_type)
    for collection_name in target_collections:
        _embed_chunks_to_collection(
            client=client,
            collection_name=collection_name,
            import_id=import_id,
            chunks=chunks,
            safe_metadata=safe_metadata,
        )


def _extract_words(text: str, max_words: int) -> str:
    words = text.split()
    return " ".join(words[:max_words])


def _parse_rows(filename: str, content: bytes) -> list[dict[str, Any]]:
    lower = filename.lower()
    if lower.endswith(".csv"):
        text = content.decode("utf-8-sig", errors="ignore")
        return list(csv.DictReader(io.StringIO(text)))
    if lower.endswith((".xlsx", ".xlsm", ".xls")):
        from openpyxl import load_workbook

        wb = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
        ws = wb.active
        rows = ws.iter_rows(values_only=True)
        headers = [str(h) if h is not None else "" for h in next(rows, [])]
        out: list[dict[str, Any]] = []
        for row in rows:
            out.append({headers[i]: row[i] for i in range(min(len(headers), len(row)))})
        return out
    return []


def extract_text_content(filename: str, file_bytes: bytes) -> str:
    lower = filename.lower()
    if lower.endswith((".txt", ".json", ".md", ".yaml", ".yml", ".log")):
        return file_bytes.decode("utf-8", errors="ignore")
    if lower.endswith((".csv", ".xlsx", ".xlsm", ".xls")):
        rows = _parse_rows(filename, file_bytes)
        headers = list(rows[0].keys()) if rows else []
        sample_lines = [" | ".join(str(row.get(header, "")) for header in headers) for row in rows[:10]]
        return f"Headers: {headers}\nSample rows:\n" + "\n".join(sample_lines)
    if lower.endswith(".pdf"):
        try:
            from pypdf import PdfReader

            reader = PdfReader(io.BytesIO(file_bytes))
            text = "\n".join((page.extract_text() or "") for page in reader.pages)
            return _extract_words(text, 500)
        except Exception:  # noqa: BLE001
            return f"PDF import received for {filename}; text extraction unavailable."
    if lower.endswith(".docx"):
        try:
            from docx import Document

            doc = Document(io.BytesIO(file_bytes))
            text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
            return _extract_words(text, 500)
        except Exception:  # noqa: BLE001
            return f"DOCX import received for {filename}; text extraction unavailable."
    return f"Binary import received for {filename}. Parse details from attached file metadata."


def extract_full_text_content(filename: str, file_bytes: bytes) -> str:
    lower = filename.lower()
    if lower.endswith((".txt", ".json", ".md", ".yaml", ".yml", ".log")):
        return file_bytes.decode("utf-8", errors="ignore")
    if lower.endswith(".csv"):
        return file_bytes.decode("utf-8-sig", errors="ignore")
    if lower.endswith((".xlsx", ".xlsm", ".xls")):
        rows = _parse_rows(filename, file_bytes)
        if not rows:
            return ""
        headers = list(rows[0].keys())
        lines = [", ".join(headers)]
        for row in rows:
            lines.append(" | ".join(str(row.get(header, "")) for header in headers))
        return "\n".join(lines)
    if lower.endswith(".pdf"):
        try:
            from pypdf import PdfReader

            reader = PdfReader(io.BytesIO(file_bytes))
            return "\n".join((page.extract_text() or "") for page in reader.pages)
        except Exception:  # noqa: BLE001
            return ""
    if lower.endswith(".docx"):
        try:
            from docx import Document

            doc = Document(io.BytesIO(file_bytes))
            return "\n".join(paragraph.text for paragraph in doc.paragraphs)
        except Exception:  # noqa: BLE001
            return ""
    return extract_text_content(filename, file_bytes)


async def _read_import_bytes(minio_path: str | None) -> bytes:
    if not minio_path:
        return b""
    client = get_minio_client()
    response = client.get_object(settings.minio_bucket, minio_path)
    payload = response.read()
    response.close()
    response.release_conn()
    return payload


def _column_value(row: dict[str, Any], column_name: str | None) -> str | None:
    if not column_name:
        return None
    if column_name in row and row[column_name] is not None:
        value = str(row[column_name]).strip()
        return value if value else None
    for key, val in row.items():
        if key.strip().lower() == column_name.strip().lower() and val is not None:
            text = str(val).strip()
            return text if text else None
    return None


def _normalize_field_mapping(column_mapping: dict[str, Any]) -> dict[str, str | None]:
    resolved: dict[str, str | None] = {
        "full_name": None,
        "email": None,
        "employee_id": None,
        "active_status": None,
        "department": None,
        "entity_company": None,
    }
    aliases = {
        "full_name": {"full_name", "name", "display_name", "employee_name", "staff_member"},
        "email": {"email", "email_address", "upn", "user_principal_name"},
        "employee_id": {"employee_id", "id", "employee_number"},
        "active_status": {"active_status", "status", "account_status", "account_enabled"},
        "department": {"department", "dept"},
        "entity_company": {"entity_company", "entity", "company", "organization"},
    }
    for canonical, keys in aliases.items():
        for key in keys:
            if key in column_mapping and column_mapping[key]:
                resolved[canonical] = str(column_mapping[key])
                break
    return resolved


def _confidence_from_analysis(detected_type: str, column_mapping: dict[str, str | None], controls: list[str]) -> str:
    mapped_count = len([value for value in column_mapping.values() if value])
    if detected_type == "unknown":
        return "low"
    if mapped_count >= 3 and controls:
        return "high"
    if mapped_count >= 1:
        return "medium"
    return "low"


async def _analyze_with_claude(sample_text: str, *, bypass_limit: bool = False) -> dict[str, Any]:
    api_key = (settings.anthropic_api_key or "").strip()
    if not api_key:
        return {
            "detected_type": "unknown",
            "column_mapping": {},
            "relevant_controls": [],
            "recommended_action": "Anthropic API key not configured; analysis unavailable.",
        }
    prompt = (
        "You are a compliance data analyst for Apprio Inc. Examine this file sample and determine:\n\n"
        "1) What type of compliance-relevant data does this file contain? Choose the best match from: "
        "active_employee_list, terminated_employee_list, mfa_enrollment, training_completion, entra_id_export, "
        "intune_compliance, crowdstrike_inventory, access_rights_report, evidence_document, policy_document, "
        "audit_log, auditor_checklist, or unknown.\n"
        "2) What are the exact column names in this file that correspond to: full name, email, employee ID, "
        "active/status, department, entity/company?\n"
        "3) Which ISO 27001 or CMMC controls is this data relevant to?\n"
        "4) What should be done with this data in a compliance context?\n\n"
        "Respond in JSON only with keys: detected_type, column_mapping (object mapping our field names to actual "
        "column names found), relevant_controls (list), recommended_action (string).\n\n"
        f"FILE SAMPLE:\n{sample_text}"
    )
    response = await call_claude(
        max_tokens=1200,
        messages=[{"role": "user", "content": prompt}],
        model=MODEL_HAIKU,
        bypass_limit=bypass_limit,
    )
    content = "".join(block.text for block in response.content if block.type == "text").strip()
    cleaned = _clean_claude_json_response(content)
    try:
        parsed = json.loads(cleaned)
    except json.JSONDecodeError:
        logger.warning("Classifier returned invalid JSON for sample; falling back to unknown type.")
        return {
            "detected_type": "unknown",
            "column_mapping": {},
            "relevant_controls": [],
            "recommended_action": "Review import and apply compliance updates.",
        }
    detected_type = str(parsed.get("detected_type", "unknown")).strip()
    if detected_type not in TYPE_CHOICES:
        detected_type = "unknown"
    return {
        "detected_type": detected_type,
        "column_mapping": parsed.get("column_mapping") or {},
        "relevant_controls": [str(item).strip() for item in parsed.get("relevant_controls") or []],
        "recommended_action": str(parsed.get("recommended_action") or ""),
    }


def _clean_claude_json_response(text: str) -> str:
    cleaned = str(text or "").strip()
    # Strip markdown code fences if present.
    if cleaned.startswith("```"):
        cleaned = cleaned.split("```", 1)[1]
        if cleaned.startswith("json"):
            cleaned = cleaned[4:]
        cleaned = cleaned.strip()
    # Also handle closing fence.
    if cleaned.endswith("```"):
        cleaned = cleaned[:-3].strip()
    return cleaned


async def _run_evidence_intelligence_with_claude(
    *,
    filename: str,
    full_text: str,
    bypass_limit: bool = False,
) -> dict[str, Any]:
    api_key = (settings.anthropic_api_key or "").strip()
    if not api_key:
        return {
            "controls": [],
            "evidence_type": "other",
            "confidence": "low",
            "summary": "Anthropic API key not configured; evidence intelligence unavailable.",
        }
    max_chars_per_chunk = 4000
    text = full_text.strip()
    chunks = [text[idx : idx + max_chars_per_chunk] for idx in range(0, len(text), max_chars_per_chunk)] or [text]
    aggregate_controls: set[str] = set()
    confidence_rank = {"low": 0, "medium": 1, "high": 2}
    best_confidence = "low"
    evidence_type_counts: dict[str, int] = {}
    summaries: list[str] = []

    for chunk_index, chunk_text in enumerate(chunks, start=1):
        prompt = (
            "You are a compliance evidence analyst for Apprio Inc. Read this document and identify every ISO 27001:2022, "
            "ISO 20000-1:2018, ISO 9001:2015, CMMC Level 2, NIST SP 800-53 Rev 5, NIST CSF 2.0, HIPAA Security Rule, and DPA Attachment C "
            "control that this document provides evidence for.\n\n"
            "Be thorough. Do not rely on the filename. Determine relevance from content alone.\n"
            "For DPA Attachment C controls, use IDs DPA.1 through DPA.23. Key mappings:\n\n"
            "Code of Conduct, anti-fraud policy, anti-bribery policy -> DPA.5, DPA.6\n"
            "Anti-fraud training records, compliance training materials -> DPA.10\n"
            "Third-party due diligence records, vendor assessments -> DPA.16, DPA.17, DPA.18\n"
            "Management commitment statements, board communications, leadership attestations -> DPA.1, DPA.2\n"
            "Ethics hotline, anonymous reporting, whistleblower protection -> DPA.12\n"
            "Investigation records, misconduct allegations, case logs -> DPA.13\n"
            "Root cause analysis, remediation documentation -> DPA.23\n"
            "Disciplinary procedures and records -> DPA.15\n"
            "Risk assessment, bribery risk register -> DPA.3, DPA.4\n"
            "Internal controls, financial controls, approval authority -> DPA.7\n"
            "Compliance program testing, internal audit of compliance program -> DPA.21, DPA.22\n"
            "Compliance officer appointment, board reporting, CCO authority -> DPA.9\n"
            "Compliance guidance system, advice channel -> DPA.11\n"
            "Anti-fraud contract provisions, third-party agreements -> DPA.18\n"
            "M&A due diligence procedures -> DPA.19, DPA.20\n"
            "Compensation and bonus criteria linked to compliance -> DPA.14\n"
            "Policy review records, program update documentation -> DPA.8\n\n"
            "For NIST SP 800-53 Rev 5 controls, use the standard control IDs (AC-1, AC-2, AC-2(1), SI-4, SC-7 etc.).\n"
            "Key mappings:\n\n"
            "MFA reports, authentication config, Entra ID conditional access -> IA-2, IA-2(1), IA-2(2), IA-5, IA-5(1)\n"
            "User access reviews, account management records -> AC-2, AC-2(3), AC-2(12)\n"
            "Offboarding/termination records -> AC-2(3), PS-4, PS-4(2)\n"
            "Audit logs (CloudTrail, Entra ID, Intune, Azure) -> AU-2, AU-3, AU-6, AU-12, AU-12(1)\n"
            "Vulnerability scan reports, patch compliance -> RA-5, RA-5(2), SI-2, SI-2(2)\n"
            "Penetration test reports -> CA-8, CA-8(1), RA-5\n"
            "CrowdStrike/EDR records -> SI-3, SI-3(1), SI-4, SI-4(2)\n"
            "Network security (VPC, firewall, security groups) -> SC-7, SC-7(3), SC-7(5), SC-7(7)\n"
            "Encryption at rest records -> SC-28, SC-28(1)\n"
            "Encryption in transit/TLS certificates -> SC-8, SC-8(1), SC-13\n"
            "Backup and recovery records -> CP-9, CP-9(1), CP-10\n"
            "Business continuity/DR plans -> CP-2, CP-4, CP-7\n"
            "Physical security site assessment -> PE-2, PE-3, PE-6, PE-8\n"
            "Training completion records -> AT-2, AT-2(2), AT-3, AT-4\n"
            "Incident response records -> IR-4, IR-4(1), IR-6, IR-8\n"
            "Risk register and risk assessment -> RA-3, RA-3(1), RA-7\n"
            "Configuration management records -> CM-2, CM-6, CM-7, CM-8\n"
            "Change management logs -> CM-3, CM-3(2)\n"
            "Asset inventory -> CM-8, CM-8(1), PM-5\n"
            "Supplier/vendor security records -> SA-9, SR-6, SR-3\n"
            "Password policy, account lockout -> IA-5(1), AC-7\n"
            "Segregation of duties -> AC-5, AC-6\n"
            "Security policies -> AC-1, AT-1, AU-1, CA-1, CM-1, CP-1, IA-1, IR-1, MA-1, MP-1, PE-1, PL-1, PS-1, RA-1, SA-1, SC-1, SI-1\n"
            "Privacy/PII/PHI records -> PT-2, PT-3, PT-5, SI-19\n"
            "System security plan, architecture docs -> PL-2, PL-8, SA-17\n"
            "Management review, board records -> PM-2, PM-9\n\n"
            "For NIST CSF 2.0, use IDs like CSF.GV.OC, CSF.ID.AM, CSF.PR.AA etc.\n"
            "For HIPAA Security Rule, use IDs like HIPAA.AS.1, HIPAA.PS.1, HIPAA.TS.1 etc.\n"
            "Return JSON only:\n"
            "{\n"
            '  "controls": ["A.6.3", "DPA.10", "DPA.1", "HIPAA.AS.2"],\n'
            '  "evidence_type": "policy|record|log|report|attestation|screenshot|config",\n'
            '  "confidence": "high|medium|low",\n'
            '  "summary": "one sentence describing what this document is"\n'
            "}\n\n"
            f"Filename: {filename}\n"
            f"Document segment {chunk_index} of {len(chunks)}:\n{chunk_text}"
        )
        parsed: dict[str, Any] | None = None
        for attempt in range(4):
            try:
                response = await asyncio.wait_for(
                    call_claude(
                        max_tokens=4096,
                        messages=[{"role": "user", "content": prompt}],
                        model=MODEL_HAIKU,
                        bypass_limit=bypass_limit,
                    ),
                    timeout=120,
                )
                response_text = "".join(block.text for block in response.content if block.type == "text").strip()
                cleaned = _clean_claude_json_response(response_text)
                try:
                    parsed = json.loads(cleaned)
                except json.JSONDecodeError:
                    logger.warning(
                        "Evidence intelligence returned invalid JSON for {} chunk {}/{}; using fallback.",
                        filename,
                        chunk_index,
                        len(chunks),
                    )
                    parsed = {
                        "controls": [],
                        "evidence_type": "other",
                        "confidence": "low",
                        "summary": f"Compliance evidence document ({filename})",
                    }
                break
            except Exception as exc:  # noqa: BLE001
                if attempt == 3:
                    raise
                backoff = 5 * (attempt + 1)
                if _is_rate_limit_error_message(str(exc)) or "Payload too large" in str(exc):
                    await asyncio.sleep(backoff)
                    continue
                await asyncio.sleep(backoff)
        if parsed is None:
            continue
        controls = [str(value).strip() for value in (parsed.get("controls") or []) if str(value).strip()]
        for control in controls:
            aggregate_controls.add(control)
        confidence = str(parsed.get("confidence") or "low").strip().lower()
        if confidence not in confidence_rank:
            confidence = "low"
        if confidence_rank[confidence] > confidence_rank[best_confidence]:
            best_confidence = confidence
        evidence_type = str(parsed.get("evidence_type") or "other").strip().lower()
        if evidence_type not in {"policy", "record", "log", "report", "attestation", "screenshot", "config"}:
            evidence_type = "other"
        evidence_type_counts[evidence_type] = evidence_type_counts.get(evidence_type, 0) + 1
        summary = str(parsed.get("summary") or "").strip()
        if summary:
            summaries.append(summary)
        await asyncio.sleep(1)

    selected_evidence_type = "other"
    if evidence_type_counts:
        selected_evidence_type = sorted(
            evidence_type_counts.items(),
            key=lambda item: item[1],
            reverse=True,
        )[0][0]
    merged_summary = summaries[0] if summaries else "Compliance evidence document."
    return {
        "controls": sorted(aggregate_controls),
        "evidence_type": selected_evidence_type,
        "confidence": best_confidence,
        "summary": merged_summary,
    }


def _map_evidence_type(raw: str, detected_type: str) -> EvidenceType:
    value = (raw or "").strip().lower()
    by_value = {
        "policy": EvidenceType.POLICY,
        "record": EvidenceType.RECORD,
        "log": EvidenceType.LOG,
        "report": EvidenceType.REPORT,
        "attestation": EvidenceType.ATTESTATION,
        "screenshot": EvidenceType.SCREENSHOT,
        "config": EvidenceType.CONFIG,
    }
    if value in by_value:
        return by_value[value]
    return (
        EvidenceType.POLICY
        if detected_type == "policy_document"
        else EvidenceType.REPORT
        if detected_type in {"intune_compliance", "audit_log", "access_rights_report", "crowdstrike_inventory"}
        else EvidenceType.RECORD
        if detected_type in {"evidence_document"}
        else EvidenceType.OTHER
    )


def _is_rate_limit_error_message(message: str | None) -> bool:
    text = (message or "").lower()
    return "429" in text or "rate limit" in text or "rate_limit" in text or "ratelimit" in text


def compute_content_hash(file_bytes: bytes) -> str:
    return hashlib.sha256(file_bytes).hexdigest()


def _duplicate_type_bucket(filename: str) -> str:
    suffix = Path(filename or "").suffix.lower()
    if suffix in _IMAGE_EXTENSIONS:
        return "image"
    if suffix in _SPREADSHEET_EXTENSIONS:
        return "spreadsheet"
    if suffix in _DOCUMENT_EXTENSIONS:
        return "document"
    if suffix in _STRUCTURED_DATA_EXTENSIONS:
        return "structured"
    return "other"


def _duplicate_similarity_tokens(text: str) -> set[str]:
    tokens = re.findall(r"[a-z0-9]{3,}", str(text or "").lower())
    return {token for token in tokens if token not in _DUPLICATE_TOKEN_STOPWORDS}


def _is_duplicate_candidate_similar(
    *,
    new_filename: str,
    new_summary: str,
    candidate_filename: str,
    candidate_summary: str,
) -> bool:
    if _duplicate_type_bucket(new_filename) != _duplicate_type_bucket(candidate_filename):
        return False

    new_name_tokens = _duplicate_similarity_tokens(Path(new_filename).stem)
    candidate_name_tokens = _duplicate_similarity_tokens(Path(candidate_filename).stem)
    name_overlap = len(new_name_tokens.intersection(candidate_name_tokens))

    new_summary_tokens = _duplicate_similarity_tokens(new_summary)
    candidate_summary_tokens = _duplicate_similarity_tokens(candidate_summary)
    summary_overlap = len(new_summary_tokens.intersection(candidate_summary_tokens))

    name_similarity = difflib.SequenceMatcher(
        None,
        Path(new_filename).stem.lower(),
        Path(candidate_filename).stem.lower(),
    ).ratio()

    return bool(name_overlap >= 2 or summary_overlap >= 3 or name_similarity >= 0.82)


async def _apply_exact_duplicate_flags(
    session: AsyncSession,
    *,
    record: DataImport,
) -> bool:
    if not record.content_hash:
        return False
    duplicate = (
        await session.execute(
            select(DataImport)
            .where(
                DataImport.content_hash == record.content_hash,
                DataImport.id != record.id,
                DataImport.library == record.library,
            )
            .order_by(DataImport.created_at.desc())
            .limit(1)
        )
    ).scalars().first()
    if duplicate is None:
        if record.duplicate_status != "false_positive":
            record.duplicate_status = "unique"
            record.duplicate_of_id = None
            record.duplicate_confidence = None
            record.duplicate_reason = None
        return False
    canonical = duplicate if duplicate.id < record.id else record
    counterpart = record if canonical.id == duplicate.id else duplicate
    canonical.duplicate_status = "confirmed_duplicate"
    canonical.duplicate_of_id = counterpart.id
    canonical.duplicate_confidence = "high"
    canonical.duplicate_reason = "Exact file content match (SHA-256)."
    canonical.duplicate_flag_dismissed = False
    counterpart.duplicate_status = "confirmed_duplicate"
    counterpart.duplicate_of_id = canonical.id
    counterpart.duplicate_confidence = "high"
    counterpart.duplicate_reason = "Exact file content match (SHA-256)."
    counterpart.duplicate_flag_dismissed = False
    return True


async def _run_fuzzy_duplicate_detection(
    session: AsyncSession,
    *,
    record: DataImport,
    bypass_limit: bool = False,
) -> bool:
    if not (record.filename or "").strip():
        return False
    summary = str(record.identified_summary or "").strip()
    if not summary:
        return False
    rows = list(
        (
            await session.execute(
                select(DataImport)
                .where(
                    DataImport.id != record.id,
                    DataImport.library == record.library,
                )
                .order_by(DataImport.created_at.desc())
                .limit(20)
            )
        ).scalars()
    )
    if not rows:
        return False
    filtered_rows = [
        row
        for row in rows
        if _is_duplicate_candidate_similar(
            new_filename=record.filename,
            new_summary=summary,
            candidate_filename=row.filename or "",
            candidate_summary=str(row.identified_summary or ""),
        )
    ]
    if not filtered_rows:
        return False
    candidates = "\n".join(
        f"- {row.id} | {row.filename} | {str(row.identified_summary or '')[:280]}"
        for row in filtered_rows
    )
    prompt = (
        "Review this document against the existing documents listed and identify ONLY cases "
        "where the documents are clearly the same document — meaning they cover the same "
        "subject matter, same time period, and same data, just saved under a different filename. "
        "Do NOT flag documents as duplicates based on file format alone, processing errors, or "
        "superficial similarities.\n"
        "A NinjaOne device inventory is never a duplicate of a customer engagement email even if both are PNGs. "
        "An IAM Users export is never a duplicate of an IAM Roles export even if both are JSON files.\n"
        "Only return suspected duplicates with HIGH confidence where the content is genuinely the same. "
        "If in doubt, return an empty list.\n\n"
        "New document:\n\n"
        f"Filename: {record.filename}\n"
        f"Summary: {summary}\n\n"
        "Existing documents:\n"
        f"{candidates}\n\n"
        "Return JSON only — empty results array if no genuine duplicates found:\n"
        "{ \"suspected_duplicates\": [] }"
    )
    try:
        response = await call_claude(
            max_tokens=800,
            messages=[{"role": "user", "content": prompt}],
            model=MODEL_HAIKU,
            bypass_limit=bypass_limit,
        )
    except Exception as exc:  # noqa: BLE001
        if is_daily_limit_exception(exc):
            logger.warning("Skipping fuzzy duplicate detection for {}: daily API limit reached", record.filename)
            return False
        raise
    text = "".join(block.text for block in response.content if block.type == "text").strip()
    try:
        payload = json.loads(_clean_claude_json_response(text))
    except json.JSONDecodeError:
        logger.warning("Skipping fuzzy duplicate detection for {}: invalid model JSON", record.filename)
        return False
    suspected = payload.get("suspected_duplicates") or []
    if not isinstance(suspected, list):
        return False
    chosen = None
    for item in suspected:
        if not isinstance(item, dict):
            continue
        confidence = str(item.get("confidence") or "").strip().lower()
        if confidence == "high":
            chosen = item
            break
    if chosen is None:
        if record.duplicate_status != "confirmed_duplicate" and record.duplicate_status != "false_positive":
            record.duplicate_status = "unique"
            record.duplicate_of_id = None
            record.duplicate_confidence = None
            record.duplicate_reason = None
        return False
    existing_import_id = chosen.get("existing_import_id")
    if not isinstance(existing_import_id, int):
        return False
    record.duplicate_status = "suspected"
    record.duplicate_of_id = existing_import_id
    record.duplicate_confidence = str(chosen.get("confidence") or "").strip().lower() or "high"
    record.duplicate_reason = str(chosen.get("reason") or "").strip()[:1000] or "Potential duplicate detected."
    record.duplicate_flag_dismissed = False
    return True


async def run_duplicate_detection_for_import(
    session: AsyncSession,
    *,
    record: DataImport,
    is_new_import: bool,
    bypass_limit: bool = False,
) -> dict[str, bool]:
    exact = await _apply_exact_duplicate_flags(session, record=record)
    fuzzy = False
    if is_new_import and not exact:
        fuzzy = await _run_fuzzy_duplicate_detection(session, record=record, bypass_limit=bypass_limit)
    if exact or fuzzy:
        await log_change(
            session,
            category="document",
            action="Duplicate flagged",
            subject=record.filename,
            detail=(
                f"Duplicate detected: {record.filename} may duplicate import #{record.duplicate_of_id}."
                if record.duplicate_of_id
                else f"Duplicate detected for {record.filename}."
            ),
        )
    return {"exact": exact, "fuzzy": fuzzy}


async def _get_control_match_index(session: AsyncSession) -> dict[str, Any]:
    cached = session.info.get("control_match_index")
    if cached is not None:
        return cached
    controls = list((await session.execute(select(Control))).scalars())
    by_id = {str(control.control_id).strip().lower(): control for control in controls}
    by_title = {str(control.title or "").strip().lower(): control for control in controls}
    index = {"controls": controls, "by_id": by_id, "by_title": by_title}
    session.info["control_match_index"] = index
    return index


def _extract_control_tokens(raw: str) -> list[str]:
    text = str(raw or "").strip()
    if not text:
        return []
    tokens: list[str] = []
    for pattern in _CONTROL_ID_PATTERNS:
        for match in re.finditer(pattern, text, flags=re.IGNORECASE):
            tokens.append(match.group(1).strip())
    clause_matches = re.findall(r"Clause\s+(\d+\.\d+)", text, flags=re.IGNORECASE)
    for value in clause_matches:
        tokens.append(f"Cl.{value}")
    return tokens


def _extract_title_phrase(raw: str) -> str | None:
    text = str(raw or "").strip()
    if not text:
        return None
    if " - " in text:
        phrase = text.split(" - ", 1)[1].strip().lower()
        return phrase if len(phrase) >= 4 else None
    return None


async def _match_controls_from_raw(
    session: AsyncSession,
    raw_controls: list[str],
) -> tuple[list[Control], list[str]]:
    if not raw_controls:
        return [], []
    index = await _get_control_match_index(session)
    by_id: dict[str, Control] = index["by_id"]
    by_title: dict[str, Control] = index["by_title"]
    controls: list[Control] = index["controls"]

    matched: dict[int, Control] = {}
    unmatched: list[str] = []

    for raw in raw_controls:
        text = str(raw or "").strip()
        if not text:
            continue
        key = text.lower()
        control = by_id.get(key)
        if control is None:
            for token in _extract_control_tokens(text):
                control = by_id.get(token.lower())
                if control is not None:
                    break
        if control is None:
            for control_id_key, maybe_control in by_id.items():
                if control_id_key in key:
                    control = maybe_control
                    break
        if control is None:
            phrase = _extract_title_phrase(text)
            if phrase:
                exact_title = by_title.get(phrase)
                if exact_title is not None:
                    control = exact_title
                else:
                    for maybe_title, maybe_control in by_title.items():
                        if phrase in maybe_title or maybe_title in phrase:
                            control = maybe_control
                            break
        if control is None:
            unmatched.append(text)
            continue
        matched[control.id] = control

    if unmatched:
        logger.warning("Unmatched controls from AI classification: {}", unmatched[:20])
    return list(matched.values()), unmatched


async def _normalize_relevant_control_ids(session: AsyncSession, ai_controls: list[str]) -> list[str]:
    matched_controls, _ = await _match_controls_from_raw(session, ai_controls)
    return sorted({control.control_id for control in matched_controls})


async def normalize_control_ids_for_matching(session: AsyncSession, raw_controls: list[str]) -> list[str]:
    matched_controls, _ = await _match_controls_from_raw(session, raw_controls)
    return sorted({control.control_id for control in matched_controls})


async def _mark_import_failed(import_id: int, reason: str) -> None:
    async with AsyncSessionLocal() as session:
        record = (await session.execute(select(DataImport).where(DataImport.id == import_id))).scalars().first()
        if record is None:
            return
        record.status = ImportStatus.FAILED
        record.error_message = reason
        record.updated_at = datetime.now(timezone.utc)
        await log_change(
            session,
            category="document",
            action="Import failed",
            subject=record.filename,
            detail=f"Import failed (import_id={record.id}): {record.filename} — {reason}",
        )
        await session.commit()


async def _run_import_with_timeout(
    import_id: int,
    *,
    content: str,
    trigger_auditor_refresh: bool,
    bypass_limit: bool = False,
) -> bool:
    try:
        await asyncio.wait_for(
            process_import_with_options(
                import_id=import_id,
                content=content,
                trigger_auditor_refresh=trigger_auditor_refresh,
                bypass_limit=bypass_limit,
            ),
            timeout=PROCESS_TIMEOUT_SECONDS,
        )
        return True
    except TimeoutError:
        await _mark_import_failed(import_id, "Processing timeout — file took too long to classify")
        return False


async def _infer_framework_labels(session: AsyncSession, control_ids: list[str]) -> str | None:
    if not control_ids:
        return None
    controls = list(
        (
            await session.execute(
                select(Control).options(selectinload(Control.framework)).where(Control.control_id.in_(control_ids))
            )
        ).scalars()
    )
    frameworks = sorted({control.framework.short_name for control in controls if control.framework})
    return ",".join(frameworks) if frameworks else None


def _is_truthy(value: str | None) -> bool:
    return (value or "").strip().lower() in {"true", "1", "yes", "enabled", "active", "complete", "completed"}


def _normalize_lookup_value(value: str | None) -> str | None:
    if not value:
        return None
    normalized = value.strip().lower()
    return normalized or None


def _coalesce_personnel_fields(target: PersonnelRecord, source: PersonnelRecord) -> None:
    str_fields = [
        "display_name",
        "email",
        "entra_upn",
        "employee_id",
        "entity",
        "termination_date",
        "training_date",
        "nda_date",
        "last_synced",
    ]
    bool_fields = [
        "active",
        "entra_account_active",
        "mfa_configured",
        "training_complete",
        "nda_on_file",
        "background_check",
    ]
    for field in str_fields:
        if getattr(target, field) in (None, "") and getattr(source, field) not in (None, ""):
            setattr(target, field, getattr(source, field))
    for field in bool_fields:
        if getattr(target, field) is None and getattr(source, field) is not None:
            setattr(target, field, getattr(source, field))
    source_flags = source.flags or []
    if source_flags:
        target_flags = target.flags or []
        target.flags = sorted({*target_flags, *source_flags})


async def _dedupe_personnel_records_for_identity(
    session: AsyncSession,
    *,
    normalized_email: str | None,
    normalized_entra_upn: str | None,
    employee_id: str | None,
) -> PersonnelRecord | None:
    clauses = []
    if normalized_email:
        clauses.append(func.lower(PersonnelRecord.email) == normalized_email)
    if normalized_entra_upn:
        clauses.append(func.lower(PersonnelRecord.entra_upn) == normalized_entra_upn)
    if employee_id:
        clauses.append(PersonnelRecord.employee_id == employee_id)
    if not clauses:
        return None

    matches = list(
        (
            await session.execute(
                select(PersonnelRecord)
                .where(or_(*clauses))
                .order_by(PersonnelRecord.id.desc())
            )
        ).scalars()
    )
    if not matches:
        return None

    keeper = matches[0]
    for duplicate in matches[1:]:
        _coalesce_personnel_fields(keeper, duplicate)
        await session.delete(duplicate)
    return keeper


async def _upsert_personnel_from_analysis(
    session: AsyncSession,
    record: DataImport,
    rows: list[dict[str, Any]],
    detected_type: str,
    mapping: dict[str, str | None],
) -> list[str]:
    updates = 0
    for row in rows[:5000]:
        full_name = _column_value(row, mapping.get("full_name"))
        email = _column_value(row, mapping.get("email"))
        employee_id = _column_value(row, mapping.get("employee_id"))
        status_text = _column_value(row, mapping.get("active_status"))
        department = _column_value(row, mapping.get("department"))
        entity = _column_value(row, mapping.get("entity_company")) or ("Canaide" if (department or "").lower().find("canaide") >= 0 else "Apprio")

        normalized_email = _normalize_lookup_value(email)
        record_row = await _dedupe_personnel_records_for_identity(
            session,
            normalized_email=normalized_email,
            normalized_entra_upn=normalized_email,
            employee_id=employee_id,
        )
        if record_row is None:
            lookup = None
            if detected_type in {"entra_id_export", "mfa_enrollment"} and email:
                lookup = await session.execute(
                    select(PersonnelRecord)
                    .where(
                        or_(
                            func.lower(PersonnelRecord.entra_upn) == normalized_email,
                            func.lower(PersonnelRecord.email) == normalized_email,
                        )
                    )
                    .order_by(PersonnelRecord.id.desc())
                )
            elif detected_type in {"active_employee_list", "terminated_employee_list"} and employee_id:
                lookup = await session.execute(
                    select(PersonnelRecord)
                    .where(PersonnelRecord.employee_id == employee_id)
                    .order_by(PersonnelRecord.id.desc())
                )
            elif detected_type == "training_completion" and full_name:
                lookup = await session.execute(
                    select(PersonnelRecord)
                    .where(PersonnelRecord.display_name.ilike(full_name))
                    .order_by(PersonnelRecord.id.desc())
                )
            else:
                lookup = await session.execute(
                    select(PersonnelRecord)
                    .where(PersonnelRecord.display_name == (full_name or "Unknown User"))
                    .order_by(PersonnelRecord.id.desc())
                )
            record_row = lookup.scalars().first() if lookup else None
        if record_row is None:
            record_row = PersonnelRecord(display_name=full_name or email or employee_id or "Unknown User", active=True)
            session.add(record_row)
        record_row.display_name = full_name or record_row.display_name
        record_row.email = email or record_row.email
        record_row.entra_upn = email or record_row.entra_upn
        record_row.employee_id = employee_id or record_row.employee_id
        record_row.entity = entity or record_row.entity

        if detected_type == "mfa_enrollment":
            record_row.mfa_configured = _is_truthy(status_text)
        elif detected_type == "training_completion":
            score = _column_value(row, "Score")
            numeric_score = None
            if score:
                try:
                    numeric_score = float(score.replace("%", "").strip())
                except ValueError:
                    numeric_score = None
            record_row.training_complete = (numeric_score is not None and numeric_score >= 80) or _is_truthy(status_text)
            completion_date = _column_value(row, "Completion Date") or _column_value(row, "Date")
            if completion_date:
                record_row.training_date = completion_date
        elif detected_type == "entra_id_export":
            record_row.entra_account_active = _is_truthy(status_text)
        elif detected_type == "terminated_employee_list":
            record_row.active = False
            if not record_row.termination_date:
                record_row.termination_date = record.data_date
        elif detected_type == "active_employee_list":
            record_row.active = True
        updates += 1
    await session.commit()
    return [f"Updated personnel records: {updates}"]


def _extract_checklist_items(rows: list[dict[str, Any]], content: str) -> list[dict[str, str]]:
    items: list[dict[str, str]] = []
    if rows:
        headers = list(rows[0].keys())
        for idx, row in enumerate(rows, start=1):
            values = [str(row.get(h, "")).strip() for h in headers]
            joined = " | ".join(v for v in values if v)
            if not joined:
                continue
            item_number = str(values[0]).strip() if values and values[0] else str(idx)
            items.append({"item_number": item_number, "description": joined})
        if items:
            return items


def _extract_dpa_control_from_text(value: str | None) -> str | None:
    text = str(value or "").strip()
    if not text:
        return None
    match = _DPA_REQUEST_ID_PATTERN.search(text)
    if not match:
        return None
    number = int(match.group(1))
    if number < 1 or number > 23:
        return None
    return f"DPA.{number}"


def _extract_request_id_from_item_payload(parsed: dict[str, Any], item: dict[str, str]) -> str:
    raw_fields = parsed.get("raw_fields") if isinstance(parsed.get("raw_fields"), dict) else {}
    request_id = str(parsed.get("reference_id") or "").strip()
    if request_id:
        return request_id
    if isinstance(raw_fields, dict):
        for key, value in raw_fields.items():
            key_text = str(key or "").strip().lower().replace(" ", "").replace("_", "")
            if key_text in {"requestid", "request", "requestnumber", "rid"}:
                candidate = str(value or "").strip()
                if candidate:
                    return candidate
    item_number = str(item.get("item_number") or "").strip()
    if item_number:
        return item_number
    return str(item.get("description") or "").strip()
    lines = [line.strip("-* \t") for line in content.splitlines() if line.strip()]
    for idx, line in enumerate(lines, start=1):
        if len(line) < 5:
            continue
        prefix = line.split(" ", 1)[0]
        looks_numbered = any(ch.isdigit() for ch in prefix)
        item_number = prefix if looks_numbered else str(idx)
        description = line if looks_numbered else line
        items.append({"item_number": item_number, "description": description})
    return items[:200]


def _build_auditor_document_sample(filename: str, rows: list[dict[str, Any]], fallback_content: str) -> str:
    if rows:
        headers = list(rows[0].keys())
        sample_lines: list[str] = []
        for idx, row in enumerate(rows[:100], start=1):
            pairs = [f"{header}={row.get(header, '')}" for header in headers]
            sample_lines.append(f"row_{idx}: " + " | ".join(pairs))
        return (
            f"FILENAME: {filename}\n"
            f"HEADERS ({len(headers)}): {headers}\n"
            "ROWS (key=value pairs):\n"
            + "\n".join(sample_lines)
        )
    return f"FILENAME: {filename}\nRAW CONTENT SAMPLE:\n{fallback_content[:30000]}"


async def _map_item_controls_with_claude(
    items: list[dict[str, str]],
    *,
    bypass_limit: bool = False,
) -> list[list[str]]:
    if not items:
        return []
    api_key = (settings.anthropic_api_key or "").strip()
    if not api_key:
        return [[] for _ in items]
    prompt = (
        "Map each auditor checklist item to relevant ISO 27001 or CMMC control IDs. "
        "Return JSON object with key 'mappings' containing an array of arrays, same order as items.\n"
        f"Items:\n{json.dumps(items)}"
    )
    try:
        response = await call_claude(
            max_tokens=1000,
            messages=[{"role": "user", "content": prompt}],
            model=MODEL_HAIKU,
            bypass_limit=bypass_limit,
        )
        text = "".join(block.text for block in response.content if block.type == "text").strip()
        payload = json.loads(_clean_claude_json_response(text))
        mappings = payload.get("mappings") or []
        normalized: list[list[str]] = []
        for mapping in mappings[: len(items)]:
            normalized.append([str(control).strip() for control in mapping if str(control).strip()])
        while len(normalized) < len(items):
            normalized.append([])
        return normalized
    except Exception:  # noqa: BLE001
        return [[] for _ in items]


async def _parse_auditor_document_with_claude(
    sample_text: str,
    *,
    bypass_limit: bool = False,
) -> dict[str, Any]:
    api_key = (settings.anthropic_api_key or "").strip()
    if not api_key:
        return {"audit_metadata": {}, "fields_found": [], "items": []}
    prompt = (
        "You are parsing an auditor's information request document for a compliance\n"
        "program. This document may be a CSV, Excel spreadsheet, PDF, Word doc, or any\n"
        "other format. Your job is to extract every distinct information request or\n"
        "checklist item the auditor is asking for.\n\n"
        "For each item found, extract:\n"
        "- A unique identifier or reference number if one exists in the document\n"
        "- A short title or category if one exists\n"
        "- The full description of what the auditor is requesting\n"
        "- Any document count targets or compliance metrics if present\n"
        "- Any status indicators if present (e.g. submitted, pending, not submitted)\n"
        "- Any notes, warnings, or flags from the auditor\n"
        "- Any deadline or due date if present\n\n"
        "Look carefully at every column in the document. Even if columns appear numeric\n"
        "or administrative, extract them. A reference or ID number is always present if\n"
        "rows are numbered. A category or topic label is present if rows have a short\n"
        "title alongside a longer description. Administrative columns showing counts,\n"
        "compliance metrics, or status flags must all be captured in raw_fields even if\n"
        "they are zero. Do not return only the description — return every field present.\n\n"
        "Respond in JSON only with this structure:\n"
        "{\n"
        "  fields_found: [list of field names actually present in the document],\n"
        "  items: [\n"
        "    {\n"
        "      reference_id: string or null,\n"
        "      title: string or null,\n"
        "      description: string (the full request text),\n"
        "      target_document_count: number or null,\n"
        "      compliant_count: number or null,\n"
        "      pending_count: number or null,\n"
        "      noncompliant_count: number or null,\n"
        "      status_from_auditor: string or null,\n"
        "      auditor_notes: string or null,\n"
        "      due_date: string or null,\n"
        "      raw_fields: object (all original field values for this item)\n"
        "    }\n"
        "  ]\n"
        "}\n\n"
        "Extract every item. Do not summarize or combine items. Preserve the auditor's exact language in the description field.\n\n"
        f"DOCUMENT SAMPLE:\n{sample_text}"
    )
    response = await call_claude(
        max_tokens=6000,
        messages=[{"role": "user", "content": prompt}],
        model=MODEL_HAIKU,
        bypass_limit=bypass_limit,
    )
    text = "".join(block.text for block in response.content if block.type == "text").strip()
    logger.info("Claude raw auditor parse response: {}", text)
    text = _clean_claude_json_response(text)
    try:
        parsed = json.loads(text)
        parsed["_raw_claude_response"] = text
        return parsed
    except json.JSONDecodeError:
        return {"fields_found": [], "items": [], "_raw_claude_response": text}


async def _create_auditor_checklist(
    session: AsyncSession,
    record: DataImport,
    rows: list[dict[str, Any]],
    content: str,
    suggested_controls: list[str],
    *,
    bypass_limit: bool = False,
) -> list[str]:
    from models.auditor import (
        AuditorChecklist,
        AuditorChecklistItem,
        AuditorChecklistStatus,
        AuditorItemPriority,
        AuditorItemStatus,
    )

    sample_text = _build_auditor_document_sample(record.filename, rows, content)
    structured = await _parse_auditor_document_with_claude(sample_text, bypass_limit=bypass_limit)
    metadata: dict[str, Any] = {}
    fields_found = [str(field) for field in (structured.get("fields_found") or [])]
    parsed_items = structured.get("items") if isinstance(structured.get("items"), list) else []
    items: list[dict[str, str]] = []
    normalized_item_payloads: list[dict[str, Any]] = []
    for idx, parsed in enumerate(parsed_items, start=1):
        if not isinstance(parsed, dict):
            continue
        description = str(parsed.get("description") or "").strip()
        if not description:
            continue
        ref = parsed.get("reference_id")
        title = parsed.get("title")
        item_number = str(ref or title or idx)
        title_prefix = f"{title}: " if title else ""
        items.append({"item_number": item_number, "description": f"{title_prefix}{description}"})
        normalized_item_payloads.append(parsed)
    if not items:
        if rows:
            headers = [str(key) for key in rows[0].keys()]
            fields_found = fields_found or headers
            ref_key = next((h for h in headers if any(token in h.lower() for token in ["rid", "reference", "req", "request id", "id#"])), None)
            title_key = next((h for h in headers if any(token in h.lower() for token in ["checklist item", "category", "title", "topic"])), None)
            desc_key = next((h for h in headers if any(token in h.lower() for token in ["information requested", "description", "request"])), None)
            for idx, row in enumerate(rows, start=1):
                raw_fields = {str(key): row.get(key) for key in headers}
                reference = str(row.get(ref_key, "")).strip() if ref_key else ""
                title = str(row.get(title_key, "")).strip() if title_key else ""
                description_value = str(row.get(desc_key, "")).strip() if desc_key else ""
                if not description_value:
                    description_value = " | ".join(
                        f"{key}={row.get(key)}" for key in headers if row.get(key) not in (None, "")
                    )
                if not description_value:
                    continue
                item_number = reference or str(idx)
                full_description = f"{title}: {description_value}" if title else description_value
                items.append({"item_number": item_number, "description": full_description})
                normalized_item_payloads.append(
                    {
                        "reference_id": reference or None,
                        "title": title or None,
                        "description": description_value,
                        "raw_fields": raw_fields,
                        "auditor_notes": str(row.get("Alerts or Warnings") or row.get("alerts_or_warnings") or "").strip() or None,
                    }
                )
        else:
            fallback = _extract_checklist_items(rows, content)
            items = [{"item_number": item["item_number"], "description": item["description"]} for item in fallback]
            normalized_item_payloads = [{"raw_fields": {}} for _ in fallback]
            if not fields_found:
                fields_found = ["description"]

    mapped_controls = await _map_item_controls_with_claude(items, bypass_limit=bypass_limit)
    engagement_name = (record.auditor_engagement_name or "").strip()
    engagement_type = (record.auditor_engagement_type or "").strip()
    engagement_year = (record.auditor_period_year or "").strip() or str(record.data_date)[:4]
    checklist_name = engagement_name or Path(record.filename).stem.replace("_", " ").replace("-", " ").strip() or "Auditor Checklist"
    checklist_audit_date = f"{engagement_year}-01-01" if engagement_year else str(record.data_date)
    checklist_auditor_name = str(record.auditor_certification_body or record.source_system or "External Auditor")
    checklist_framework = str(record.framework or ("dpa_attachment_c" if (record.library or "main") == "dpa" else "mixed"))
    existing_checklists = list(
        (
            await session.execute(
                select(AuditorChecklist).where(AuditorChecklist.status == AuditorChecklistStatus.ACTIVE)
            )
        ).scalars()
    )
    checklist = None
    for existing in existing_checklists:
        if not record.auditor_merge_with_existing:
            continue
        if (existing.name or "").strip() != checklist_name:
            continue
        if (existing.audit_type or "").strip() != engagement_type:
            continue
        if (existing.audit_period_year or "").strip() != engagement_year:
            continue
        checklist = existing
        break

    merged = checklist is not None
    if checklist is None:
        checklist = AuditorChecklist(
            name=checklist_name,
            audit_type=engagement_type,
            audit_period_year=engagement_year,
            audit_date=checklist_audit_date,
            auditor_name=checklist_auditor_name,
            framework=checklist_framework,
            status=AuditorChecklistStatus.ACTIVE,
            source_import_id=record.id,
            fields_found=fields_found,
        )
        session.add(checklist)
        await session.flush()
    else:
        checklist.name = checklist_name
        checklist.audit_type = checklist.audit_type or engagement_type
        checklist.audit_period_year = checklist.audit_period_year or engagement_year
        checklist.audit_date = checklist.audit_date or checklist_audit_date
        checklist.auditor_name = checklist.auditor_name or checklist_auditor_name
        checklist.framework = checklist.framework or checklist_framework
        checklist.fields_found = sorted({*(checklist.fields_found or []), *fields_found})

    existing_signatures: set[str] = set()
    if merged:
        existing_items = list(
            (
                await session.execute(
                    select(AuditorChecklistItem).where(AuditorChecklistItem.checklist_id == checklist.id)
                )
            ).scalars()
        )
        for existing_item in existing_items:
            existing_signatures.add(
                f"{(existing_item.item_number or '').strip().lower()}::{(existing_item.description or '').strip().lower()}"
            )
    inserted_count = 0
    for idx, item in enumerate(items):
        parsed = normalized_item_payloads[idx] if idx < len(normalized_item_payloads) else {}
        request_id = _extract_request_id_from_item_payload(parsed, item) if isinstance(parsed, dict) else str(item.get("item_number") or "")
        dpa_control_id = _extract_dpa_control_from_text(request_id)
        if dpa_control_id is None:
            dpa_control_id = _extract_dpa_control_from_text(item.get("description"))
        if dpa_control_id is None:
            dpa_control_id = _extract_dpa_control_from_text(item.get("item_number"))
        if dpa_control_id:
            control_ids = [dpa_control_id]
        else:
            control_ids = mapped_controls[idx] if idx < len(mapped_controls) else []
        due_date = parsed.get("due_date") if isinstance(parsed, dict) else None
        due_date_str = str(due_date) if due_date else None
        auditor_notes = str(parsed.get("auditor_notes") or "") if isinstance(parsed, dict) else ""
        raw_fields = parsed.get("raw_fields") if isinstance(parsed, dict) and isinstance(parsed.get("raw_fields"), dict) else {}
        signature = f"{(item['item_number'] or '').strip().lower()}::{(item['description'] or '').strip().lower()}"
        if signature in existing_signatures:
            continue
        session.add(
            AuditorChecklistItem(
                checklist_id=checklist.id,
                source_import_id=record.id,
                item_number=item["item_number"] or str(idx + 1),
                description=item["description"],
                control_ids=control_ids,
                status=AuditorItemStatus.OPEN,
                our_response="",
                evidence_ids=[],
                due_date=due_date_str,
                auditor_notes=auditor_notes or None,
                priority=AuditorItemPriority.MEDIUM,
                raw_fields=raw_fields,
                evidence_mapping={"results": []},
            )
        )
        inserted_count += 1
    await session.commit()
    from services.auditor_evidence_mapper import trigger_mapping_job

    trigger_mapping_job(checklist.id)
    return [
        f"{'Merged into' if merged else 'Created'} auditor checklist #{checklist.id} with {inserted_count} new items.",
        f"Fields found: {', '.join(fields_found) if fields_found else 'none'}",
    ]


async def reparse_checklist_control_mappings(
    session: AsyncSession,
    checklist_id: int,
    *,
    bypass_limit: bool = False,
) -> dict[str, int]:
    from models.auditor import AuditorChecklist, AuditorChecklistItem

    checklist = (
        await session.execute(select(AuditorChecklist).where(AuditorChecklist.id == checklist_id))
    ).scalars().first()
    if checklist is None:
        return {"checklist_id": checklist_id, "items_total": 0, "items_updated": 0}

    items = list(
        (
            await session.execute(
                select(AuditorChecklistItem)
                .where(AuditorChecklistItem.checklist_id == checklist_id)
                .order_by(AuditorChecklistItem.id.asc())
            )
        ).scalars()
    )
    if not items:
        return {"checklist_id": checklist_id, "items_total": 0, "items_updated": 0}

    map_input = [
        {
            "item_number": str(item.item_number or ""),
            "description": str(item.description or ""),
        }
        for item in items
    ]
    mapped_controls = await _map_item_controls_with_claude(map_input, bypass_limit=bypass_limit)

    updated = 0
    for idx, item in enumerate(items):
        raw_fields = item.raw_fields if isinstance(item.raw_fields, dict) else {}
        request_hint = str(
            raw_fields.get("Request_ID")
            or raw_fields.get("request_id")
            or raw_fields.get("Request ID")
            or item.item_number
            or item.description
            or ""
        )
        dpa_control_id = _extract_dpa_control_from_text(request_hint)
        if dpa_control_id is None:
            dpa_control_id = _extract_dpa_control_from_text(item.description)
        if dpa_control_id is None:
            dpa_control_id = _extract_dpa_control_from_text(item.item_number)

        if dpa_control_id:
            new_control_ids = [dpa_control_id]
        else:
            new_control_ids = mapped_controls[idx] if idx < len(mapped_controls) else []

        existing = [str(value) for value in (item.control_ids or [])]
        if existing != new_control_ids:
            item.control_ids = new_control_ids
            updated += 1

    await session.commit()
    return {
        "checklist_id": checklist_id,
        "items_total": len(items),
        "items_updated": updated,
    }


async def _upsert_evidence_with_controls(
    session: AsyncSession,
    *,
    filename: str,
    minio_path: str | None,
    data_date: str,
    source_system: str,
    detected_type: str,
    evidence_type: EvidenceType,
    relevant_controls: list[str],
    recommended_action: str,
    confidence: str | None = None,
    summary: str | None = None,
    library: str = "main",
) -> dict[str, int | bool]:
    matched_controls, unmatched_controls = await _match_controls_from_raw(session, relevant_controls)
    controls = sorted(matched_controls, key=lambda row: row.control_id)
    evidence = (
        await session.execute(
            select(EvidenceItem)
            .options(selectinload(EvidenceItem.controls))
            .where(
                func.lower(EvidenceItem.filename) == filename.strip().lower(),
                EvidenceItem.library == library,
            )
        )
    ).scalars().first()
    created = False
    if evidence is None:
        evidence = EvidenceItem(
            filename=filename,
            file_path=minio_path,
            evidence_type=evidence_type,
            description=f"{detected_type.replace('_', ' ').title()} import",
            entity=source_system,
            collected_date=data_date,
            status=EvidenceStatus.CURRENT,
            notes=recommended_action or f"Imported from {source_system}",
            analysis_confidence=confidence,
            analysis_summary=summary,
            library=library,
        )
        session.add(evidence)
        created = True
    else:
        evidence.file_path = minio_path or evidence.file_path
        evidence.evidence_type = evidence_type
        evidence.description = evidence.description or f"{detected_type.replace('_', ' ').title()} import"
        evidence.entity = evidence.entity or source_system
        evidence.collected_date = evidence.collected_date or data_date
        evidence.status = EvidenceStatus.CURRENT
        if recommended_action:
            evidence.notes = recommended_action
        if confidence:
            evidence.analysis_confidence = confidence
        if summary:
            evidence.analysis_summary = summary
        evidence.library = library
    if summary:
        evidence.description = summary

    existing_controls = {control.id: control for control in (evidence.controls or [])}
    linked_controls = 0
    for control in controls:
        if control.id not in existing_controls:
            linked_controls += 1
        existing_controls[control.id] = control
    evidence.controls = list(existing_controls.values())

    evidenced_controls = 0
    for control in controls:
        if control.status != ControlStatus.EVIDENCED:
            control.status = ControlStatus.EVIDENCED
        evidenced_controls += 1
    await session.commit()
    return {
        "created": created,
        "linked_controls": linked_controls,
        "evidenced_controls": evidenced_controls,
        "unmatched_controls": len(unmatched_controls),
    }


async def _store_as_evidence(
    session: AsyncSession,
    record: DataImport,
    detected_type: str,
    relevant_controls: list[str],
    recommended_action: str,
    row_count: int,
) -> list[str]:
    result = await _upsert_evidence_with_controls(
        session,
        filename=record.filename,
        minio_path=record.minio_path,
        data_date=record.data_date,
        source_system=record.source_system,
        detected_type=detected_type,
        evidence_type=_map_evidence_type("", detected_type),
        relevant_controls=relevant_controls,
        recommended_action=recommended_action,
        library=record.library or "main",
    )
    return [
        f"{'Created' if result['created'] else 'Updated'} evidence record and linked controls.",
        f"Rows/doc units processed: {row_count}",
        f"Controls evidenced: {result['evidenced_controls']}",
        f"Unmatched controls: {result['unmatched_controls']}",
    ]


async def _run_evidence_intelligence_pass(
    session: AsyncSession,
    *,
    record: DataImport,
    full_text: str,
    detected_type: str,
    bypass_limit: bool = False,
) -> dict[str, Any]:
    if not full_text.strip():
        return {
            "controls": [],
            "confidence": "low",
            "summary": "No extractable text content.",
            "evidence_type": _map_evidence_type("", detected_type),
            "links_created": 0,
            "unmatched_controls": 0,
        }
    intelligence = await _run_evidence_intelligence_with_claude(
        filename=record.filename,
        full_text=full_text,
        bypass_limit=bypass_limit,
    )
    normalized_controls = await _normalize_relevant_control_ids(
        session,
        [str(value) for value in intelligence.get("controls") or []],
    )
    evidence_type = _map_evidence_type(str(intelligence.get("evidence_type") or ""), detected_type)
    summary = str(intelligence.get("summary") or "").strip()
    confidence = str(intelligence.get("confidence") or "low").strip().lower()
    upsert_result = await _upsert_evidence_with_controls(
        session,
        filename=record.filename,
        minio_path=record.minio_path,
        data_date=record.data_date,
        source_system=record.source_system,
        detected_type=detected_type,
        evidence_type=evidence_type,
        relevant_controls=normalized_controls,
        recommended_action=summary or "Evidence intelligence analysis complete.",
        confidence=confidence,
        summary=summary,
        library=record.library or "main",
    )
    return {
        "controls": normalized_controls,
        "confidence": confidence,
        "summary": summary,
        "evidence_type": evidence_type.value,
        "links_created": int(upsert_result["linked_controls"]),
        "unmatched_controls": int(upsert_result["unmatched_controls"]),
    }


async def run_evidence_intelligence_on_import(
    record: DataImport,
    session: AsyncSession,
    *,
    bypass_limit: bool = False,
    max_chars: int | None = None,
) -> dict[str, Any]:
    if _is_blocked_import_filename(record.filename or ""):
        raise ValueError("Blocked filename for import analysis")
    payload = await _read_import_bytes(record.minio_path)
    full_text = extract_full_text_content(record.filename, payload)
    if not full_text.strip():
        full_text = extract_text_content(record.filename, payload)
    if max_chars is not None and max_chars > 0 and len(full_text) > max_chars:
        full_text = full_text[:max_chars]
    detected_type = "policy_document" if "policy" in record.filename.lower() else "evidence_document"
    intelligence = await _run_evidence_intelligence_pass(
        session,
        record=record,
        full_text=full_text,
        detected_type=detected_type,
        bypass_limit=bypass_limit,
    )
    controls = list(intelligence.get("controls") or [])
    if controls:
        if not record.control_ids:
            record.control_ids = controls
        else:
            record.control_ids = sorted(set(record.control_ids or []).union(controls))
    record.updated_at = datetime.now(timezone.utc)
    await session.commit()
    return intelligence


async def backfill_evidence_links(
    session: AsyncSession,
    *,
    limit: int | None = None,
) -> dict[str, int]:
    stmt = select(DataImport).where(DataImport.status == ImportStatus.COMPLETE).order_by(DataImport.id.asc())
    if limit is not None:
        stmt = stmt.limit(limit)
    imports = list((await session.execute(stmt)).scalars())
    stats = {
        "total_complete_imports": len(imports),
        "processed": 0,
        "links_created": 0,
        "imports_skipped_no_controls": 0,
        "unmatched_controls": 0,
        "errors": 0,
    }
    for record in imports:
        raw_controls = record.control_ids or []
        if not raw_controls:
            stats["imports_skipped_no_controls"] += 1
            continue
        if isinstance(raw_controls, str):
            try:
                raw_controls = json.loads(raw_controls)
            except json.JSONDecodeError:
                raw_controls = []
        if not raw_controls:
            stats["imports_skipped_no_controls"] += 1
            continue
        stats["processed"] += 1
        detected_type = "policy_document" if "policy" in record.filename.lower() else "evidence_document"
        try:
            result = await _upsert_evidence_with_controls(
                session,
                filename=record.filename,
                minio_path=record.minio_path,
                data_date=record.data_date,
                source_system=record.source_system,
                detected_type=detected_type,
                evidence_type=_map_evidence_type("", detected_type),
                relevant_controls=[str(value) for value in raw_controls],
                recommended_action=record.identified_summary or "Backfilled evidence linkage from completed import.",
                library=record.library or "main",
            )
            stats["links_created"] += int(result["linked_controls"])
            stats["unmatched_controls"] += int(result["unmatched_controls"])
        except Exception as exc:  # noqa: BLE001
            stats["errors"] += 1
            logger.error("Backfill failed for import {} ({}): {}", record.id, record.filename, exc)
    return stats


async def reanalyze_all_evidence_imports(
    session: AsyncSession,
    *,
    limit: int | None = None,
    progress_callback: Callable[[dict[str, int]], Awaitable[None]] | None = None,
    start_from: int = 0,
    bypass_limit: bool = False,
) -> dict[str, int]:
    checkpoint = await _get_reanalyze_checkpoint(session)
    effective_start = max(int(start_from or 0), checkpoint)
    stmt = (
        select(DataImport)
        .where(
            DataImport.status == ImportStatus.COMPLETE,
            DataImport.id > effective_start,
        )
        .order_by(DataImport.id.asc())
    )
    if limit is not None:
        stmt = stmt.limit(limit)
    imports = list((await session.execute(stmt)).scalars())
    logger.info(
        "Starting evidence reanalysis: start_from={} checkpoint={} effective_start={} total={}",
        start_from,
        checkpoint,
        effective_start,
        len(imports),
    )
    stats = {
        "total_complete_imports": len(imports),
        "reanalyzed": 0,
        "reembedded": 0,
        "new_control_links": 0,
        "skipped_already_analyzed": 0,
        "errors": 0,
    }
    processed = 0
    last_processed_import_id = effective_start
    for record in imports:
        try:
            logger.info(
                "Reanalyzing import_id={} filename='{}' ({}/{})",
                record.id,
                record.filename,
                processed + 1,
                stats["total_complete_imports"],
            )
            existing_evidence = (
                await session.execute(
                    select(EvidenceItem)
                    .options(selectinload(EvidenceItem.controls))
                    .where(func.lower(EvidenceItem.filename) == record.filename.strip().lower())
                )
            ).scalars().first()
            if (
                existing_evidence is not None
                and (existing_evidence.analysis_summary or "").strip()
                and existing_evidence.controls
            ):
                stats["skipped_already_analyzed"] += 1
                processed += 1
                last_processed_import_id = record.id
                if progress_callback is not None:
                    await progress_callback(
                        {
                            "completed": processed,
                            "total": stats["total_complete_imports"],
                            "new_links": stats["new_control_links"],
                            "errors": stats["errors"],
                        }
                    )
                continue
            payload = await _read_import_bytes(record.minio_path)
            full_text = extract_full_text_content(record.filename, payload)
            if not full_text.strip():
                full_text = extract_text_content(record.filename, payload)
            detected_type = "policy_document" if "policy" in record.filename.lower() else "evidence_document"
            embedding_text = full_text
            if len(embedding_text) > 200000:
                embedding_text = embedding_text[:200000]
            if full_text.strip():
                embed_import_text(
                    import_id=record.id,
                    text=embedding_text,
                    metadata={
                        "filename": record.filename,
                        "source_system": record.source_system,
                        "data_date": record.data_date,
                        "framework": record.framework or "",
                    },
                    detected_type=detected_type,
                )
                stats["reembedded"] += 1
            intelligence = await _run_evidence_intelligence_pass(
                session,
                record=record,
                full_text=full_text,
                detected_type=detected_type,
                bypass_limit=bypass_limit,
            )
            if full_text.strip():
                embed_import_text(
                    import_id=record.id,
                    text=embedding_text,
                    metadata={
                        "filename": record.filename,
                        "source_system": record.source_system,
                        "data_date": record.data_date,
                        "framework": record.framework or "",
                        "analysis_summary": str(intelligence.get("summary") or ""),
                    },
                    detected_type=detected_type,
                )
            controls = list(intelligence.get("controls") or [])
            if controls:
                if not record.control_ids:
                    record.control_ids = controls
                else:
                    record.control_ids = sorted(set(record.control_ids or []).union(controls))
            stats["new_control_links"] += int(intelligence.get("links_created") or 0)
            record.updated_at = datetime.now(timezone.utc)
            await session.commit()
            stats["reanalyzed"] += 1
            last_processed_import_id = record.id
        except Exception as exc:  # noqa: BLE001
            stats["errors"] += 1
            logger.error("Reanalyze failed for import {} ({}): {}", record.id, record.filename, exc)
            last_processed_import_id = record.id
        processed += 1
        if processed % 10 == 0 and last_processed_import_id > 0:
            await _set_reanalyze_checkpoint(session, last_processed_import_id)
            logger.info(
                "Reanalysis checkpoint saved at import_id={} (processed {}/{})",
                last_processed_import_id,
                processed,
                stats["total_complete_imports"],
            )
        if progress_callback is not None:
            await progress_callback(
                {
                    "completed": processed,
                    "total": stats["total_complete_imports"],
                    "new_links": stats["new_control_links"],
                    "errors": stats["errors"],
                }
            )
    if stats["errors"] == 0:
        await _clear_reanalyze_checkpoint(session)
        logger.info("Reanalysis completed successfully. Checkpoint cleared.")
    return stats


def _sample_for_ai(filename: str, file_bytes: bytes) -> tuple[str, list[dict[str, Any]]]:
    lower = filename.lower()
    if lower.endswith((".csv", ".xlsx", ".xlsm", ".xls")):
        rows = _parse_rows(filename, file_bytes)
        headers = list(rows[0].keys()) if rows else []
        sample_lines = []
        for row in rows[:10]:
            pairs = [f"{header}={row.get(header, '')}" for header in headers]
            sample_lines.append(" | ".join(pairs))
        return f"Headers: {headers}\nRows:\n" + "\n".join(sample_lines), rows
    extracted = extract_text_content(filename, file_bytes)
    return _extract_words(extracted, 500), []


def _override_from_notes(notes: str | None) -> str | None:
    if not notes:
        return None
    marker = "[override_type:"
    if marker not in notes:
        return None
    start = notes.index(marker) + len(marker)
    end = notes.find("]", start)
    if end == -1:
        return None
    value = notes[start:end].strip()
    return value if value in TYPE_CHOICES else None


def _infer_detected_type_from_record(record: DataImport) -> str:
    for entry in record.proposed_updates or []:
        if not isinstance(entry, str):
            continue
        if not entry.lower().startswith("detected type:"):
            continue
        detected = entry.split(":", 1)[1].strip().lower()
        if detected:
            return detected
    override = _override_from_notes(record.notes)
    if override:
        return override
    filename_lower = (record.filename or "").lower()
    if any(token in filename_lower for token in ["meeting minutes", "meeting_notes", "meeting-notes", "minutes"]):
        return "meeting_notes"
    return "policy_document" if "policy" in filename_lower else "evidence_document"


def _has_compliance_embedding_for_filename(filename: str) -> bool:
    client = chromadb.HttpClient(host=settings.chroma_host, port=settings.chroma_port)
    collection = client.get_or_create_collection("compliance_docs")
    existing = collection.get(where={"filename": filename}, limit=1)
    return bool(existing.get("ids"))


async def fix_compliance_doc_embeddings(
    session: AsyncSession,
    *,
    threshold: int | None = None,
) -> dict[str, int]:
    client = chromadb.HttpClient(host=settings.chroma_host, port=settings.chroma_port)
    compliance = client.get_or_create_collection("compliance_docs")
    current_count = int(compliance.count())
    if threshold is not None and current_count >= threshold:
        return {
            "skipped": 1,
            "current_count": current_count,
            "checked_imports": 0,
            "reembedded": 0,
            "missing_text": 0,
            "errors": 0,
            "final_count": current_count,
        }

    imports = list(
        (
            await session.execute(
                select(DataImport)
                .where(DataImport.status == ImportStatus.COMPLETE)
                .order_by(DataImport.id.asc())
            )
        ).scalars()
    )
    stats = {
        "skipped": 0,
        "current_count": current_count,
        "checked_imports": len(imports),
        "reembedded": 0,
        "missing_text": 0,
        "errors": 0,
        "final_count": current_count,
    }
    for record in imports:
        try:
            if _has_compliance_embedding_for_filename(record.filename):
                continue
            payload = await _read_import_bytes(record.minio_path)
            full_text = extract_full_text_content(record.filename, payload)
            if not full_text.strip():
                full_text = extract_text_content(record.filename, payload)
            if not full_text.strip():
                stats["missing_text"] += 1
                continue
            embedding_text = full_text[:200000] if len(full_text) > 200000 else full_text
            detected_type = _infer_detected_type_from_record(record)
            embed_import_text(
                import_id=record.id,
                text=embedding_text,
                metadata={
                    "filename": record.filename,
                    "source_system": record.source_system,
                    "data_date": record.data_date,
                    "framework": record.framework or "",
                    "analysis_summary": record.identified_summary or "",
                },
                detected_type=detected_type,
                force_collections=("compliance_docs",),
            )
            stats["reembedded"] += 1
        except Exception as exc:  # noqa: BLE001
            stats["errors"] += 1
            logger.error("Embedding repair failed for import {} ({}): {}", record.id, record.filename, exc)
    stats["final_count"] = int(compliance.count())
    return stats


async def run_embedding_migration_if_needed(*, threshold: int = 400) -> dict[str, int]:
    async with AsyncSessionLocal() as session:
        return await fix_compliance_doc_embeddings(session, threshold=threshold)


async def backfill_missing_content_hashes_if_needed(
    *,
    null_ratio_threshold: float = 0.5,
) -> dict[str, int | float | bool]:
    async with AsyncSessionLocal() as session:
        total_imports = int((await session.execute(select(func.count(DataImport.id)))).scalar() or 0)
        if total_imports == 0:
            return {
                "ran": False,
                "total_imports": 0,
                "null_hash_imports": 0,
                "null_ratio": 0.0,
                "updated": 0,
                "errors": 0,
            }
        null_hash_imports = int(
            (
                await session.execute(
                    select(func.count(DataImport.id)).where(
                        or_(DataImport.content_hash.is_(None), DataImport.content_hash == "")
                    )
                )
            ).scalar()
            or 0
        )
        null_ratio = float(null_hash_imports) / float(total_imports)
        if null_ratio <= null_ratio_threshold:
            return {
                "ran": False,
                "total_imports": total_imports,
                "null_hash_imports": null_hash_imports,
                "null_ratio": null_ratio,
                "updated": 0,
                "errors": 0,
            }

        imports_to_backfill = list(
            (
                await session.execute(
                    select(DataImport).where(
                        or_(DataImport.content_hash.is_(None), DataImport.content_hash == "")
                    )
                )
            ).scalars()
        )

        updated = 0
        errors = 0
        minio_client = get_minio_client()
        for row in imports_to_backfill:
            if not row.minio_path:
                errors += 1
                continue
            try:
                response = minio_client.get_object(settings.minio_bucket, row.minio_path)
                payload = response.read()
                response.close()
                response.release_conn()
                row.content_hash = compute_content_hash(payload)
                if row.file_size is None:
                    row.file_size = len(payload)
                row.updated_at = datetime.now(timezone.utc)
                updated += 1
            except Exception:  # noqa: BLE001
                errors += 1
        await session.commit()
        return {
            "ran": True,
            "total_imports": total_imports,
            "null_hash_imports": null_hash_imports,
            "null_ratio": null_ratio,
            "updated": updated,
            "errors": errors,
        }


async def process_import(import_id: int, content: str) -> None:
    await _run_import_with_timeout(
        import_id=import_id,
        content=content,
        trigger_auditor_refresh=True,
        bypass_limit=False,
    )


async def _trigger_active_checklist_mapping(session: AsyncSession) -> None:
    from models.auditor import AuditorChecklist, AuditorChecklistStatus
    from services.auditor_evidence_mapper import trigger_mapping_job

    active_checklists = list(
        (
            await session.execute(
                select(AuditorChecklist).where(
                    AuditorChecklist.status == AuditorChecklistStatus.ACTIVE
                )
            )
        ).scalars()
    )
    for checklist in active_checklists:
        trigger_mapping_job(checklist.id)


async def process_import_with_options(
    import_id: int,
    content: str,
    *,
    trigger_auditor_refresh: bool = True,
    bypass_limit: bool = False,
    is_new_import: bool = True,
) -> None:
    async with AsyncSessionLocal() as session:
        record = (await session.execute(select(DataImport).where(DataImport.id == import_id))).scalars().first()
        if record is None:
            return
        if _is_blocked_import_filename(record.filename or ""):
            record.status = ImportStatus.FAILED
            record.error_message = "Blocked filename for import processing."
            record.updated_at = datetime.now(timezone.utc)
            await log_change(
                session,
                category="document",
                action="Import failed",
                subject=record.filename,
                detail=f"Import failed (import_id={record.id}): {record.filename} — {record.error_message}",
            )
            await session.commit()
            return
        try:
            record.status = ImportStatus.PROCESSING
            await session.commit()

            file_bytes = await _read_import_bytes(record.minio_path)
            if not content.strip():
                content = extract_text_content(record.filename, file_bytes)
            full_text = extract_full_text_content(record.filename, file_bytes)
            embedding_text = full_text or content
            if len(embedding_text) > 200000:
                embedding_text = embedding_text[:200000]
            embed_import_text(
                import_id=record.id,
                text=embedding_text,
                metadata={
                    "filename": record.filename,
                    "source_system": record.source_system,
                    "data_date": record.data_date,
                    "framework": record.framework or "",
                },
                detected_type=_override_from_notes(record.notes) or _infer_detected_type_from_record(record),
            )
            ingestion = await ingest_text(session, content=content, source_label=f"{record.source_system}:{record.filename}")
            extracted_actions = (
                ingestion.get("extracted_actions", [])
                if isinstance(ingestion, dict)
                else getattr(ingestion, "extracted_actions", [])
            )
            ai_sample, rows = _sample_for_ai(record.filename, file_bytes)
            extension = Path(record.filename).suffix.lower()
            shortcut_document_extensions = {".pdf", ".docx", ".doc", ".txt", ".md", ".yaml", ".yml", ".png", ".jpg", ".jpeg"}
            should_shortcut = bool(record.control_ids) and extension in shortcut_document_extensions
            if should_shortcut:
                detected_type = (
                    "policy_document"
                    if "policy" in record.filename.lower()
                    else "evidence_document"
                )
                column_mapping = {}
                relevant_controls = sorted(set(record.control_ids or []))
                normalized_controls = await _normalize_relevant_control_ids(session, relevant_controls)
                recommended_action = "Linked document using previously identified controls."
                confidence = "high"
            else:
                analysis = await _analyze_with_claude(ai_sample, bypass_limit=bypass_limit)
                override_type = _override_from_notes(record.notes)
                detected_type = override_type or analysis["detected_type"]
                column_mapping = _normalize_field_mapping(analysis["column_mapping"])
                relevant_controls = sorted(set(analysis["relevant_controls"]))
                normalized_controls = await _normalize_relevant_control_ids(session, relevant_controls)
                recommended_action = analysis["recommended_action"] or "Review import and apply compliance updates."
                confidence = _confidence_from_analysis(detected_type, column_mapping, relevant_controls)

            if not record.control_ids:
                record.control_ids = normalized_controls
            elif normalized_controls:
                record.control_ids = sorted(set(record.control_ids or []).union(normalized_controls))
            if not record.framework:
                record.framework = await _infer_framework_labels(session, record.control_ids or [])

            updates: list[str] = []
            intelligence_result: dict[str, Any] | None = None
            if detected_type != "auditor_checklist":
                intelligence_result = await _run_evidence_intelligence_pass(
                    session,
                    record=record,
                    full_text=full_text or content,
                    detected_type=detected_type,
                    bypass_limit=bypass_limit,
                )
                intelligence_controls = list(intelligence_result.get("controls") or [])
                if intelligence_controls:
                    if not record.control_ids:
                        record.control_ids = intelligence_controls
                    else:
                        record.control_ids = sorted(set(record.control_ids or []).union(intelligence_controls))
                    if not record.framework:
                        record.framework = await _infer_framework_labels(session, record.control_ids or [])
                updates.extend(
                    [
                        f"Evidence intelligence confidence: {intelligence_result.get('confidence', 'low')}",
                        f"Evidence intelligence controls: {', '.join(intelligence_controls) if intelligence_controls else 'none'}",
                        f"Evidence intelligence links created: {intelligence_result.get('links_created', 0)}",
                    ]
                )
                embed_import_text(
                    import_id=record.id,
                    text=embedding_text,
                    metadata={
                        "filename": record.filename,
                        "source_system": record.source_system,
                        "data_date": record.data_date,
                        "framework": record.framework or "",
                        "analysis_summary": str(intelligence_result.get("summary") or ""),
                    },
                    detected_type=detected_type,
                )
            elif embedding_text.strip():
                embed_import_text(
                    import_id=record.id,
                    text=embedding_text,
                    metadata={
                        "filename": record.filename,
                        "source_system": record.source_system,
                        "data_date": record.data_date,
                        "framework": record.framework or "",
                        "analysis_summary": recommended_action,
                    },
                    detected_type=detected_type,
                )
            if detected_type in {
                "active_employee_list",
                "terminated_employee_list",
                "mfa_enrollment",
                "training_completion",
                "entra_id_export",
            }:
                updates.extend(
                    await _upsert_personnel_from_analysis(
                        session=session,
                        record=record,
                        rows=rows,
                        detected_type=detected_type,
                        mapping=column_mapping,
                    )
                )
            elif detected_type == "auditor_checklist":
                updates.extend(
                    await _create_auditor_checklist(
                        session=session,
                        record=record,
                        rows=rows,
                        content=content,
                        suggested_controls=record.control_ids or [],
                        bypass_limit=bypass_limit,
                    )
                )
            elif detected_type in {
                "intune_compliance",
                "crowdstrike_inventory",
                "access_rights_report",
                "evidence_document",
                "policy_document",
                "audit_log",
                "unknown",
            } and intelligence_result is None:
                updates.extend(
                    await _store_as_evidence(
                        session=session,
                        record=record,
                        detected_type=detected_type,
                        relevant_controls=record.control_ids or [],
                        recommended_action=recommended_action,
                        row_count=len(rows) if rows else len(ai_sample.split()),
                    )
                )

            record.status = ImportStatus.COMPLETE
            controls_text = ", ".join(record.control_ids or []) or "none"
            record.identified_summary = (
                f"{recommended_action.strip() or 'Review import and apply compliance updates.'} "
                f"Relevant controls: {controls_text}."
            ).strip()
            record.proposed_updates = [
                f"Detected type: {detected_type}",
                f"Detection confidence: {confidence}",
                f"Relevant controls (AI): {controls_text}",
                f"Recommended action: {recommended_action}",
            ] + updates + [str(action) for action in extracted_actions[:3]]
            record.updated_at = datetime.now(timezone.utc)
            await run_duplicate_detection_for_import(
                session,
                record=record,
                is_new_import=is_new_import,
                bypass_limit=bypass_limit,
            )
            action_label = "Document imported" if is_new_import else "Document updated"
            detail_label = (
                f"Document imported: {record.filename} — linked to {len(record.control_ids or [])} controls"
                if is_new_import
                else f"Document updated: {record.filename} — {len(record.control_ids or [])} control links refreshed"
            )
            await log_change(
                session,
                category="document",
                action=action_label,
                subject=record.filename,
                detail=detail_label,
            )
            await session.commit()
            if trigger_auditor_refresh:
                await _trigger_active_checklist_mapping(session)
        except Exception as exc:  # noqa: BLE001
            if is_daily_limit_exception(exc):
                record.status = ImportStatus.QUEUED
                record.error_message = (
                    "Daily API limit reached — queued until limit reset or operator override."
                )
            else:
                record.status = ImportStatus.FAILED
                record.error_message = str(exc)
                await log_change(
                    session,
                    category="document",
                    action="Import failed",
                    subject=record.filename,
                    detail=f"Import failed (import_id={record.id}): {record.filename} — {record.error_message}",
                )
            record.updated_at = datetime.now(timezone.utc)
            await session.commit()


async def process_batch_import(
    batch_id: str,
    import_ids: list[int],
    *,
    delay_between_ai_calls_seconds: float = 1.0,
    bypass_limit: bool = False,
) -> None:
    effective_delay = 2.0 if len(import_ids) > 100 else delay_between_ai_calls_seconds
    for import_id in import_ids:
        attempt = 0
        while True:
            await _run_import_with_timeout(
                import_id=import_id,
                content="",
                trigger_auditor_refresh=False,
                bypass_limit=bypass_limit,
            )

            async with AsyncSessionLocal() as session:
                record = (
                    await session.execute(select(DataImport).where(DataImport.id == import_id))
                ).scalars().first()
                if record is None:
                    break
                if record.status == ImportStatus.COMPLETE:
                    break
                if (
                    record.status == ImportStatus.QUEUED
                    and is_daily_limit_exception(record.error_message)
                ):
                    remaining_ids = [value for value in import_ids if value >= import_id]
                    for remaining_id in remaining_ids:
                        queued_record = (
                            await session.execute(select(DataImport).where(DataImport.id == remaining_id))
                        ).scalars().first()
                        if queued_record is None:
                            continue
                        if queued_record.status == ImportStatus.COMPLETE:
                            continue
                        queued_record.status = ImportStatus.QUEUED
                        queued_record.error_message = (
                            "Daily API limit reached — queued until limit reset or operator override."
                        )
                        queued_record.updated_at = datetime.now(timezone.utc)
                    await session.commit()
                    logger.warning(
                        "Paused batch {} because daily API limit was reached at import_id={}",
                        batch_id,
                        import_id,
                    )
                    return

                if attempt >= MAX_IMPORT_RETRIES:
                    record.status = ImportStatus.FAILED
                    record.error_message = "Max retries exceeded"
                    record.updated_at = datetime.now(timezone.utc)
                    await log_change(
                        session,
                        category="document",
                        action="Import failed",
                        subject=record.filename,
                        detail=f"Import failed (import_id={record.id}): {record.filename} — {record.error_message}",
                    )
                    await session.commit()
                    break

                attempt += 1
                record.retry_count = (record.retry_count or 0) + 1
                last_error = record.error_message or ""
                await session.commit()

                logger.warning(
                    "Retrying import {} ({}) attempt {}/{}",
                    record.id,
                    record.filename,
                    attempt,
                    MAX_IMPORT_RETRIES,
                )
                if _is_rate_limit_error_message(last_error):
                    logger.warning(
                        "Rate limit hit on {} — waiting 60 seconds before retry",
                        record.filename,
                    )
                    await asyncio.sleep(60)
            continue

        if effective_delay > 0:
            await asyncio.sleep(effective_delay)

    async with AsyncSessionLocal() as session:
        await _trigger_active_checklist_mapping(session)

