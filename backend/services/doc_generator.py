from __future__ import annotations

import io
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from minio import Minio
from openpyxl import Workbook
from openpyxl.styles import PatternFill
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import selectinload

from config import get_settings
from models.compliance import Control, EvidenceItem, Finding, FindingStatus, Framework, Obligation
from services.gap_scanner import scan_framework_readiness

settings = get_settings()


def _minio() -> Minio:
    return Minio(
        endpoint=settings.minio_endpoint,
        access_key=settings.minio_access_key,
        secret_key=settings.minio_secret_key,
        secure=False,
    )


def _upload_bytes(filename: str, content: bytes, content_type: str) -> tuple[str, str]:
    client = _minio()
    if not client.bucket_exists(settings.minio_bucket):
        client.make_bucket(settings.minio_bucket)
    object_name = f"generated/{filename}"
    client.put_object(
        settings.minio_bucket,
        object_name,
        io.BytesIO(content),
        len(content),
        content_type=content_type,
    )
    return object_name, f"http://localhost:8010/reports/download/{filename}"


def _doc_with_header() -> Document:
    template = Path(__file__).resolve().parents[2] / "templates" / "apprio_base.docx"
    doc = Document(str(template)) if template.exists() else Document()
    doc.core_properties.author = "Michael DuPlantis"
    doc.core_properties.last_modified_by = "Michael DuPlantis"
    header = doc.sections[0].header
    if not header.paragraphs:
        header.add_paragraph()
    header.paragraphs[0].text = "APPRIO"
    return doc


async def generate_gap_report(framework_short_name: str, session: AsyncSession) -> dict[str, str]:
    framework_result = await session.execute(
        select(Framework).where(Framework.short_name == framework_short_name)
    )
    framework = framework_result.scalar_one_or_none()
    if framework is None:
        raise ValueError("Framework not found")

    controls_result = await session.execute(
        select(Control)
        .options(selectinload(Control.evidence_items), selectinload(Control.evidence_requirements))
        .where(Control.framework_id == framework.id)
    )
    controls = list(controls_result.scalars())
    wb = Workbook()
    ws = wb.active
    ws.title = "Gaps"
    ws.append(
        [
            "Control ID",
            "Control Title",
            "Domain",
            "Status",
            "Evidence Required",
            "Evidence on File",
            "Gap Description",
            "Owner",
        ]
    )
    green = PatternFill("solid", fgColor="C6EFCE")
    yellow = PatternFill("solid", fgColor="FFEB9C")
    red = PatternFill("solid", fgColor="FFC7CE")
    domain_summary: dict[str, int] = {}

    for control in controls:
        required = ", ".join(req.evidence_type.value for req in control.evidence_requirements)
        on_file = ", ".join(item.filename for item in control.evidence_items)
        gap_desc = "Ready" if control.status.value == "evidenced" else "Needs evidence/update"
        ws.append(
            [
                control.control_id,
                control.title,
                control.domain,
                control.status.value,
                required,
                on_file,
                gap_desc,
                control.owner_default or "",
            ]
        )
        row_idx = ws.max_row
        fill = red
        if control.status.value == "evidenced":
            fill = green
        elif control.status.value == "in_progress":
            fill = yellow
        for col in range(1, 9):
            ws.cell(row=row_idx, column=col).fill = fill
        domain_summary[control.domain] = domain_summary.get(control.domain, 0) + 1

    summary_sheet = wb.create_sheet("Summary")
    summary_sheet.append(["Domain", "Control Count"])
    for domain, count in sorted(domain_summary.items()):
        summary_sheet.append([domain, count])

    buffer = io.BytesIO()
    wb.save(buffer)
    filename = f"{framework_short_name}_GapReport_{datetime.now(timezone.utc).strftime('%Y%m%d')}.xlsx"
    minio_path, url = _upload_bytes(
        filename,
        buffer.getvalue(),
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )
    return {"filename": filename, "minio_path": minio_path, "download_url": url}


async def generate_scorecard(session: AsyncSession) -> dict[str, str]:
    readiness = await scan_framework_readiness(session)
    findings_result = await session.execute(select(Finding).where(Finding.status != FindingStatus.CLOSED))
    findings = list(findings_result.scalars())
    obligations_result = await session.execute(select(Obligation))
    obligations = list(obligations_result.scalars())

    buf = io.BytesIO()
    pdf = canvas.Canvas(buf, pagesize=letter)
    pdf.setTitle("Management Scorecard")
    pdf.setFont("Helvetica-Bold", 16)
    pdf.drawString(40, 760, "Apprio Management Scorecard")
    y = 720
    pdf.setFont("Helvetica", 10)
    for row in readiness:
        pct = float(row["percent_evidenced"])
        color = colors.red if pct < 70 else colors.orange if pct < 90 else colors.green
        pdf.drawString(40, y, f"{row['name']}: {pct:.2f}%")
        pdf.setFillColor(color)
        pdf.rect(220, y - 4, min(300, pct * 3), 10, fill=1, stroke=0)
        pdf.setFillColor(colors.black)
        y -= 20

    y -= 10
    pdf.drawString(40, y, f"Open Findings: {len(findings)}")
    y -= 18
    due_soon = sum(1 for o in obligations if o.status.value in {"due_soon", "overdue"})
    pdf.drawString(40, y, f"Obligations due soon/overdue: {due_soon}")
    y -= 18
    audit_date = datetime(2026, 5, 1, tzinfo=timezone.utc)
    days = (audit_date - datetime.now(timezone.utc)).days
    pdf.drawString(40, y, f"Days to audit: {days}")
    pdf.save()

    filename = f"MGMT_Scorecard_{datetime.now(timezone.utc).strftime('%Y%m%d')}.pdf"
    minio_path, url = _upload_bytes(filename, buf.getvalue(), "application/pdf")
    return {"filename": filename, "minio_path": minio_path, "download_url": url}


async def generate_corrective_action_report(session: AsyncSession) -> dict[str, str]:
    findings_result = await session.execute(
        select(Finding).options(selectinload(Finding.corrective_actions)).where(Finding.status != FindingStatus.CLOSED)
    )
    findings = list(findings_result.scalars())
    doc = _doc_with_header()
    doc.add_heading("Corrective Action Report", level=1)
    for finding in findings:
        doc.add_heading(f"{finding.finding_id}: {finding.title}", level=2)
        doc.add_paragraph(f"Status: {finding.status.value} | Severity: {finding.severity.value}")
        doc.add_paragraph(finding.description or "")
        for action in finding.corrective_actions:
            doc.add_paragraph(
                f"- {action.description} | Owner: {action.owner or 'Unassigned'} | Due: {action.due_date or 'N/A'} | Status: {action.status.value}"
            )
    buffer = io.BytesIO()
    doc.save(buffer)
    filename = f"CA_Report_{datetime.now(timezone.utc).strftime('%Y%m%d')}.docx"
    minio_path, url = _upload_bytes(
        filename,
        buffer.getvalue(),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    return {"filename": filename, "minio_path": minio_path, "download_url": url}


async def generate_audit_package_index(framework_short_name: str, session: AsyncSession) -> dict[str, str]:
    framework_result = await session.execute(
        select(Framework).where(Framework.short_name == framework_short_name)
    )
    framework = framework_result.scalar_one_or_none()
    if framework is None:
        raise ValueError("Framework not found")

    controls_result = await session.execute(
        select(Control).options(selectinload(Control.evidence_items)).where(Control.framework_id == framework.id)
    )
    controls = list(controls_result.scalars())
    doc = _doc_with_header()
    doc.add_heading(f"Audit Package Index - {framework.name}", level=1)
    for control in sorted(controls, key=lambda c: (c.domain, c.control_id)):
        doc.add_heading(f"{control.control_id} - {control.title}", level=2)
        if not control.evidence_items:
            doc.add_paragraph("No evidence on file.")
        for evidence in control.evidence_items:
            doc.add_paragraph(
                f"{evidence.filename} | {evidence.evidence_type.value} | Collected: {evidence.collected_date or 'N/A'} | Status: {evidence.status.value}"
            )
    buffer = io.BytesIO()
    doc.save(buffer)
    filename = f"{framework_short_name}_AuditIndex_{datetime.now(timezone.utc).strftime('%Y%m%d')}.docx"
    minio_path, url = _upload_bytes(
        filename,
        buffer.getvalue(),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    return {"filename": filename, "minio_path": minio_path, "download_url": url}
